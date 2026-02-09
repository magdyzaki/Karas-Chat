# InvoicesPage.py
from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QLabel, QHBoxLayout, QLineEdit, QPushButton,
    QComboBox, QTableWidget, QTableWidgetItem, QMessageBox, QFileDialog, QInputDialog
)
from PyQt5.QtGui import QFont
from PyQt5.QtCore import Qt, QDateTime
import sqlite3, os
import locale

# محاولة استيراد مكتبة docx (اختيارية)
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    Pt = None
    Inches = None
    WD_ALIGN_PARAGRAPH = None
    OxmlElement = None
    qn = None

try:
    from num2words import num2words
except Exception:
    num2words = None

DB = os.path.join(os.path.dirname(__file__), "..", "database", "crm.db")


class InvoicesPage(QWidget):
    def __init__(self):
        super().__init__()
        # إزالة الستايل الثابت للسماح بتطبيق الستايل من MainWindow
        # self.setStyleSheet("background-color: #FFFDF5;")  # تم إزالة الستايل الثابت
        self.invoice_items = []
        self.save_folder = os.path.join(os.path.dirname(__file__), "..", "exports", "invoices")
        os.makedirs(self.save_folder, exist_ok=True)
        self.payment_terms = ""
        self.bank_details = ""
        self.seller_info = ""
        self.customer_data = {}
        self._cached_total = 0.0
        self._cached_qty = 0.0

        self.init_ui()
        self.load_customers_list()
        self.load_settings()

    def init_ui(self):
        main = QVBoxLayout()
        title = QLabel("🧾 Invoices Management")
        title.setFont(QFont("Amiri", 18, QFont.Bold))
        title.setAlignment(Qt.AlignCenter)
        main.addWidget(title)

        top_row = QHBoxLayout()
        self.invoice_number_input = QLineEdit()
        self.invoice_number_input.setPlaceholderText("Invoice No (manual)")
        self.invoice_number_input.setFixedWidth(220)
        self.save_folder_btn = QPushButton("Choose Save Folder")
        self.save_folder_btn.clicked.connect(self.choose_save_folder)
        top_row.addWidget(QLabel("Invoice No:"))
        top_row.addWidget(self.invoice_number_input)
        top_row.addStretch()
        top_row.addWidget(self.save_folder_btn)
        main.addLayout(top_row)

        # صف اختيار العميل
        customer_row = QHBoxLayout()
        self.customer_combo = QComboBox()
        self.customer_combo.setMinimumWidth(200)
        self.customer_combo.currentIndexChanged.connect(self.on_customer_changed)  # تحميل تلقائي عند التغيير
        self.load_customers_button = QPushButton("🔄 تحديث العملاء")
        self.load_customers_button.clicked.connect(self.load_customers_list)
        self.load_sales_by_customer_button = QPushButton("📋 تحميل المبيعات")
        self.load_sales_by_customer_button.clicked.connect(self.load_sales_by_customer)
        customer_row.addWidget(QLabel("اختر العميل:"))
        customer_row.addWidget(self.customer_combo)
        customer_row.addWidget(self.load_customers_button)
        customer_row.addWidget(self.load_sales_by_customer_button)
        customer_row.addStretch()
        main.addLayout(customer_row)
        
        # صف اختيار المبيعات
        import_row = QHBoxLayout()
        self.sales_combo = QComboBox()
        self.sales_combo.setMinimumWidth(300)
        self.add_sale_btn = QPushButton("Add Sale to Invoice")
        self.add_sale_btn.clicked.connect(self.add_selected_sale_to_invoice)
        import_row.addWidget(QLabel("اختر عملية البيع:"))
        import_row.addWidget(self.sales_combo)
        import_row.addWidget(self.add_sale_btn)
        import_row.addStretch()
        main.addLayout(import_row)

        self.table = QTableWidget()
        self.table.setColumnCount(8)
        self.table.setHorizontalHeaderLabels([
            "Sale ID", "Customer", "Code", "Product", "Unit", "Qty", "Price", "Total (USD)"
        ])
        self.table.horizontalHeader().setStretchLastSection(True)
        main.addWidget(self.table)

        buttons_row = QHBoxLayout()
        self.remove_item_btn = QPushButton("Remove Item")
        self.clear_items_btn = QPushButton("Clear Items")
        self.preview_btn = QPushButton("Preview Invoice")
        self.export_word_btn = QPushButton("Export Word")

        self.remove_item_btn.clicked.connect(self.remove_selected_item)
        self.clear_items_btn.clicked.connect(self.clear_items)
        self.preview_btn.clicked.connect(self.preview_invoice)
        self.export_word_btn.clicked.connect(self.export_word)

        for btn, color in [
            (self.remove_item_btn, "#E53935"),
            (self.clear_items_btn, "#9C27B0"),
            (self.preview_btn, "#03A9F4"),
            (self.export_word_btn, "#2196F3"),
        ]:
            btn.setStyleSheet(f"background-color:{color};color:white;")
            btn.setFont(QFont("Amiri", 12, QFont.Bold))
            btn.setFixedHeight(40)
            buttons_row.addWidget(btn)
        main.addLayout(buttons_row)

        summary_row = QHBoxLayout()
        self.total_label = QLabel("Total: 0.00 USD")
        self.total_label.setFont(QFont("Amiri", 14, QFont.Bold))
        summary_row.addStretch()
        summary_row.addWidget(self.total_label)
        main.addLayout(summary_row)

        extra_row = QHBoxLayout()

        self.term_combo = QComboBox()
        self.term_combo.addItems(["CIF", "CAD", "CFR"])
        extra_row.addWidget(QLabel("Term:"))
        extra_row.addWidget(self.term_combo)

        self.container_type_combo = QComboBox()
        self.container_type_combo.addItems(["20", "40", "40 HQ"])
        extra_row.addWidget(QLabel("Container Type:"))
        extra_row.addWidget(self.container_type_combo)

        self.container_count_combo = QComboBox()
        self.container_count_combo.addItems([str(i) for i in range(1, 11)])
        extra_row.addWidget(QLabel("Container Count:"))
        extra_row.addWidget(self.container_count_combo)

        main.addLayout(extra_row)

        settings_row = QHBoxLayout()

        self.payment_terms_btn = QPushButton("Payment Terms (edit)")
        self.payment_terms_btn.clicked.connect(self.edit_payment_terms)
        self.bank_details_btn = QPushButton("Bank Details (edit)")
        self.bank_details_btn.clicked.connect(self.edit_bank_details)
        self.seller_info_btn = QPushButton("Seller Info (edit)")
        self.seller_info_btn.clicked.connect(self.edit_seller_info)
        settings_row.addWidget(self.payment_terms_btn)
        settings_row.addWidget(self.bank_details_btn)
        settings_row.addWidget(self.seller_info_btn)
        main.addLayout(settings_row)

        self.setLayout(main)

    def db_conn(self):
        return sqlite3.connect(DB)

    def choose_save_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Choose Save Folder", self.save_folder)
        if folder:
            self.save_folder = folder

    def load_customers_list(self):
        """تحميل قائمة العملاء"""
        try:
            # منع الإشارة أثناء التحديث
            self.customer_combo.blockSignals(True)
            self.customer_combo.clear()
            conn = self.db_conn()
            cur = conn.cursor()
            cur.execute("SELECT DISTINCT name FROM customers ORDER BY name COLLATE NOCASE")
            rows = cur.fetchall()
            conn.close()
            if not rows:
                self.customer_combo.addItem("(لا عملاء)", None)
            else:
                self.customer_combo.addItem("(جميع العملاء)", None)
                for r in rows:
                    self.customer_combo.addItem(r[0], r[0])
            self.customer_combo.blockSignals(False)
            QMessageBox.information(self, "تم", f"تم تحميل {len(rows) if rows else 0} عميل.")
        except Exception as e:
            self.customer_combo.blockSignals(False)
            QMessageBox.critical(self, "خطأ", f"فشل تحميل العملاء:\n{str(e)}")
    
    def on_customer_changed(self, index):
        """عند تغيير العميل المحدد"""
        # لا نفعل شيئاً تلقائياً، المستخدم يضغط على "تحميل المبيعات" يدوياً
        pass
    
    def load_sales_by_customer(self):
        """تحميل المبيعات الخاصة بالعميل المحدد"""
        try:
            self.sales_combo.clear()
            customer_name = self.customer_combo.currentData()
            
            conn = self.db_conn()
            cur = conn.cursor()
            
            if customer_name is None or customer_name == "(جميع العملاء)":
                # تحميل جميع المبيعات
                cur.execute("SELECT id, customer_name, product_name, sale_date FROM sales ORDER BY id DESC")
            else:
                # تحميل مبيعات العميل المحدد
                cur.execute("""
                    SELECT id, customer_name, product_name, sale_date 
                    FROM sales 
                    WHERE customer_name = ? 
                    ORDER BY id DESC
                """, (customer_name,))
            
            rows = cur.fetchall()
            conn.close()
            
            if not rows:
                self.sales_combo.addItem("(لا توجد مبيعات)", None)
                QMessageBox.information(self, "تنبيه", "لا توجد مبيعات للعميل المحدد.")
            else:
                for r in rows:
                    display_text = f"#{r[0]} | {r[1]} | {r[2]} | {r[3]}"
                    self.sales_combo.addItem(display_text, r[0])
                QMessageBox.information(self, "تم", f"تم تحميل {len(rows)} عملية بيع.")
        except Exception as e:
            QMessageBox.critical(self, "خطأ", f"فشل تحميل المبيعات:\n{str(e)}")

    def add_selected_sale_to_invoice(self):
        company = ""
        sale_id = self.sales_combo.currentData()
        if not sale_id:
            QMessageBox.warning(self, "Warning", "Choose a sale first.")
            return
        try:
            conn = self.db_conn()
            cur = conn.cursor()
            
            # جلب بيانات البيع
            cur.execute("""
                SELECT id, customer_id, customer_name, COALESCE(product_code, ''),
                       product_name, unit, quantity, COALESCE(price_usd, 0)
                FROM sales WHERE id=?
            """, (sale_id,))
            row = cur.fetchone()

            if not row:
                conn.close()
                return

            sid, customer_id, cust_name, code, pname, unit, qty, price = row

            # جلب بيانات العميل من جدول customers
            cust_addr = '———'
            cust_phone = '———'
            company = ''
            
            if customer_id:
                try:
                    # محاولة جلب address أولاً، ثم استخدام country أو email كبديل
                    cur.execute("""
                        SELECT company, phone, COALESCE(address, country, email, '') 
                        FROM customers 
                        WHERE id = ?
                        LIMIT 1
                    """, (customer_id,))
                    c_row = cur.fetchone()
                    if c_row:
                        company = c_row[0] or ''
                        cust_phone = c_row[1] or '———'
                        cust_addr = c_row[2] or '———'
                except Exception as e:
                    # إذا فشل (مثل عدم وجود عمود address)، نجرب بدون address
                    try:
                        cur.execute("""
                            SELECT company, phone, country, email 
                            FROM customers 
                            WHERE id = ?
                            LIMIT 1
                        """, (customer_id,))
                        c_row = cur.fetchone()
                        if c_row:
                            company = c_row[0] or ''
                            cust_phone = c_row[1] or '———'
                            cust_addr = c_row[2] or c_row[3] or '———'
                    except:
                        pass
            
            # إذا لم نجد customer_id، نحاول البحث بالاسم
            if cust_addr == '———' and cust_name:
                try:
                    # محاولة جلب address أولاً، ثم استخدام country أو email كبديل
                    cur.execute("""
                        SELECT company, phone, COALESCE(address, country, email, '') 
                        FROM customers 
                        WHERE LOWER(TRIM(name)) = LOWER(TRIM(?))
                        LIMIT 1
                    """, (cust_name,))
                    c_row = cur.fetchone()
                    if c_row:
                        company = c_row[0] or ''
                        cust_phone = c_row[1] or '———'
                        cust_addr = c_row[2] or '———'
                except Exception as e:
                    # إذا فشل (مثل عدم وجود عمود address)، نجرب بدون address
                    try:
                        cur.execute("""
                            SELECT company, phone, country, email 
                            FROM customers 
                            WHERE LOWER(TRIM(name)) = LOWER(TRIM(?))
                            LIMIT 1
                        """, (cust_name,))
                        c_row = cur.fetchone()
                        if c_row:
                            company = c_row[0] or ''
                            cust_phone = c_row[1] or '———'
                            cust_addr = c_row[2] or c_row[3] or '———'
                    except:
                        pass

            conn.close()

            qty_val = float(qty or 0)
            price_val = float(price or 0)
            total = round(qty_val * price_val, 2)

            self.invoice_items.append({
                "sale_id": sid,
                "customer_name": cust_name,
                "product_code": code,
                "product_name": pname,
                "unit": unit or "",
                "quantity": qty_val,
                "price": price_val,
                "total": total
            })

            self.customer_data = {
                "name": cust_name,
                "company": company,  # مؤقتًا نفس الاسم لحد ما نضيف company في الداتا بيز
                "address": cust_addr,
                "phone": cust_phone
            }

            self.refresh_invoice_table()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Add failed:\n{e}")

    def refresh_invoice_table(self):
        self.table.setRowCount(0)
        for it in self.invoice_items:
            r = self.table.rowCount()
            self.table.insertRow(r)
            for i, val in enumerate([
                it.get("sale_id"), it.get("customer_name"), it.get("product_code"),
                it.get("product_name"), it.get("unit"), it.get("quantity"),
                self.format_money(it.get("price")), self.format_money(it.get("total"))
            ]):
                item = QTableWidgetItem(str(val))
                item.setTextAlignment(Qt.AlignCenter)
                self.table.setItem(r, i, item)
        self.update_totals()

    def remove_selected_item(self):
        try:
            row = self.table.currentRow()
            if row >= 0 and row < len(self.invoice_items):
                del self.invoice_items[row]
                self.refresh_invoice_table()
            else:
                QMessageBox.warning(self, "تنبيه", "اختر عنصراً لحذفه.")
        except Exception as e:
            QMessageBox.critical(self, "خطأ", f"حدث خطأ أثناء حذف العنصر:\n{str(e)}")

    def clear_items(self):
        self.invoice_items = []
        self.refresh_invoice_table()

    def update_totals(self):
        total = sum(it["total"] for it in self.invoice_items)
        qty_total = sum(it["quantity"] for it in self.invoice_items)
        self.total_label.setText(f"Total: {total:.2f} USD")
        self._cached_total = total
        self._cached_qty = qty_total

    def format_money(self, val):
        try:
            num = float(val)
            # لو الرقم صحيح (يعني 0 بعد العلامة العشرية)
            if num.is_integer():
                return str(int(num))
            else:
                return f"{num:.2f}"
        except:
            return str(val)

    def load_settings(self):
        try:
            conn = self.db_conn()
            cur = conn.cursor()
            keys = ["payment_terms", "bank_details", "seller_info"]
            for key in keys:
                cur.execute("SELECT value FROM settings WHERE key=?", (key,))
                row = cur.fetchone()
                setattr(self, key, row[0] if row else "")
            conn.close()
        except Exception:
            pass

    def save_setting(self, key, val):
        try:
            conn = self.db_conn()
            cur = conn.cursor()
            cur.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", (key, val))
            conn.commit()
            conn.close()
        except Exception as e:
            QMessageBox.warning(self, "Warning", f"Failed to save setting:\n{e}")

    def edit_payment_terms(self):
        txt, ok = QInputDialog.getMultiLineText(
            self, "Edit Payment Terms", "Enter payment terms:", self.payment_terms
        )
        if ok:
            self.payment_terms = txt
            self.save_setting("payment_terms", txt)

    def edit_bank_details(self):
        txt, ok = QInputDialog.getMultiLineText(
            self, "Edit Bank Details", "Enter bank details:", self.bank_details
        )
        if ok:
            self.bank_details = txt
            self.save_setting("bank_details", txt)

    def edit_seller_info(self):
        txt, ok = QInputDialog.getMultiLineText(
            self, "Edit Seller Info", "Enter seller info:", self.seller_info
        )
        if ok:
            self.seller_info = txt
            self.save_setting("seller_info", txt)

    def preview_invoice(self):
        self.generate_invoice_word(preview=True)

    def export_word(self):
        self.generate_invoice_word(preview=False)

    def generate_invoice_word(self, preview=False):
        try:
            # التحقق من توفر مكتبة docx
            if not DOCX_AVAILABLE:
                QMessageBox.warning(
                    self, 
                    "مكتبة غير مثبتة", 
                    "مكتبة python-docx غير مثبتة.\n\n"
                    "لتثبيتها، شغّل الأمر التالي في Terminal:\n"
                    "pip install python-docx\n\n"
                    "أو:\n"
                    "pip install python-docx num2words"
                )
                return
            
            if not self.invoice_items:
                QMessageBox.warning(self, "Warning", "No items to export.")
                return

            invoice_no = self.invoice_number_input.text().strip() or QDateTime.currentDateTime().toString("yyyyMMddHHmmss")
            folder = os.path.join(self.save_folder, "preview" if preview else "")
            os.makedirs(folder, exist_ok=True)
            path = os.path.join(folder, f"invoice_{invoice_no}.docx")

            doc = Document()
            style = doc.styles['Normal']
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)

            section = doc.sections[0]
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

            doc.add_paragraph("")
            doc.add_paragraph("")
            doc.add_paragraph("")

            title_p = doc.add_paragraph()
            title_run = title_p.add_run(f"INVOICE NO {invoice_no}")
            title_run.bold = True
            title_run.font.size = Pt(20)
            title_run.font.name = "Times New Roman"
            title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_p.paragraph_format.space_after = Pt(3)

            doc.add_paragraph("")
            # ===== Customer Info Section =====
            locale.setlocale(locale.LC_TIME, "en_US.UTF-8")
            current_date = QDateTime.currentDateTime().toPyDateTime().strftime("%B %d, %Y")

            # نكتب سطور العميل
            # ===== DATE =====
            p_date = doc.add_paragraph()
            r1 = p_date.add_run("DATE: ")
            r1.font.size = Pt(16)
            r1.bold = True
            r1.font.name = "Times New Roman"

            r2 = p_date.add_run(current_date)
            r2.font.size = Pt(12)
            r2.bold = True
            r2.font.name = "Times New Roman"

            # ===== TO =====
            p_to = doc.add_paragraph()
            r1 = p_to.add_run("TO: ")
            r1.font.size = Pt(16)
            r1.bold = True
            r1.font.name = "Times New Roman"

            r2 = p_to.add_run(self.customer_data.get('company') or self.customer_data.get('name', ''))
            r2.font.size = Pt(12)
            r2.bold = True
            r2.font.name = "Times New Roman"

            # ===== ADDRESS =====
            p_addr = doc.add_paragraph()
            r1 = p_addr.add_run("ADDRESS: ")
            r1.font.size = Pt(16)
            r1.bold = True
            r1.font.name = "Times New Roman"

            r2 = p_addr.add_run("ul. Rzeczna 2, 62-404 Ciążeń, Poland")
            r2.font.size = Pt(12)
            r2.bold = True
            r2.font.name = "Times New Roman"

            # ===== TEL NO =====
            p_tel = doc.add_paragraph()
            r1 = p_tel.add_run("TEL NO: ")
            r1.font.size = Pt(16)
            r1.bold = True
            r1.font.name = "Times New Roman"

            r2 = p_tel.add_run(self.customer_data.get('phone', ''))
            r2.font.size = Pt(12)
            r2.bold = True
            r2.font.name = "Times New Roman"

            doc.add_paragraph("")


            seller_lines = self.seller_info.split("\n") if self.seller_info else [
                "Seller: ELRAEE FOR DEHYDRATION",
                "Plot No. 107, Light Industry Zone, New Beni Suef, Egypt",
                "Tel: +20 122 908 4204"
            ]
            for line in seller_lines:
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(1)
                for r in p.runs:
                    r.font.size = Pt(12)
                    r.bold = True
                    r.font.name = "Times New Roman"

            doc.add_paragraph("").paragraph_format.space_after = Pt(0.1)

            table = doc.add_table(rows=1, cols=4)
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table.autofit = False

            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.append(tblPr)

            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '8000')
            tblW.set(qn('w:type'), 'dxa')
            tblPr.append(tblW)

            col_widths = [Inches(10.0), Inches(0.4), Inches(0.5), Inches(0.5)]
            for i, w in enumerate(col_widths):
                for cell in table.columns[i].cells:
                    cell.width = w

            hdrs = ["DESCRIPTION", "QTY", "PRICE", "USD"]
            hdr_row = table.rows[0].cells
            for i, h in enumerate(hdrs):
                hdr_row[i].text = h
                for p in hdr_row[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        r.bold = True
                        r.font.size = Pt(11)
                        r.font.name = "Times New Roman"

            total_value = 0
            total_qty = 0
            for it in self.invoice_items:
                row = table.add_row().cells
                row[0].text = it["product_name"]
                row[1].text = str(int(it["quantity"])) if float(it["quantity"]).is_integer() else str(it["quantity"])
                row[2].text = self.format_money(it["price"])
                row[3].text = self.format_money(it["total"])
                total_value += it["total"]
                total_qty += it["quantity"]

                for cell in row:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs:
                            r.bold = True
                            r.font.size = Pt(11)
                            r.font.name = "Times New Roman"

            total_words = num2words(int(total_value), lang='en').upper() + " DOLLARS" if num2words else f"{total_value} DOLLARS"
            qty_words = num2words(int(total_qty), lang='en').upper() + " KGS" if num2words else f"{total_qty} KGS"

            total_row = table.add_row().cells
            total_row[0].text = total_words
            total_row[1].text = "———"
            total_row[2].text = "———"
            total_row[3].text = self.format_money(total_value)

            qty_row = table.add_row().cells
            qty_row[0].text = qty_words
            qty_row[1].text = str(int(total_qty))
            qty_row[2].text = "———"
            qty_row[3].text = "———"

            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs:
                            r.bold = True
                            r.font.size = Pt(11)
                            r.font.name = "Times New Roman"

            for row in table.rows:
                for cell in row.cells:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    borders = OxmlElement('w:tcBorders')
                    for side in ['top', 'left', 'bottom', 'right']:
                        border = OxmlElement(f'w:{side}')
                        border.set(qn('w:val'), 'single')
                        border.set(qn('w:sz'), '8')
                        border.set(qn('w:color'), '000000')
                        borders.append(border)
                    tcPr.append(borders)

            try:
                normal_style = doc.styles['Normal']
                normal_pf = normal_style.paragraph_format
                normal_pf.space_before = Pt(0)
                normal_pf.space_after = Pt(0)
                normal_pf.line_spacing = 1.0
            except Exception:
                pass

            terms = doc.add_paragraph()
            terms.paragraph_format.space_before = Pt(0)
            terms.paragraph_format.space_after = Pt(0)
            terms.paragraph_format.line_spacing = 1.0

            run_title = terms.add_run("Terms of Payment ")
            run_title.bold = True
            run_title.font.name = "Times New Roman"
            run_title.font.size = Pt(14)

            run_body = terms.add_run(self.payment_terms or
                                     "70% Cash Against Documents on first presentation.\n30% after arrival and approval of goods.")
            run_body.bold = True
            run_body.font.name = "Times New Roman"
            run_body.font.size = Pt(14)

            weight_total = sum(it["quantity"] for it in self.invoice_items)
            term_value = self.term_combo.currentText()
            container_type = self.container_type_combo.currentText()
            container_count = self.container_count_combo.currentText()

            doc.add_paragraph("")
            doc.add_paragraph("")
            extra_info_p = doc.add_paragraph()
            extra_info_p.add_run(f"TERM: {term_value}\n").bold = True
            extra_info_p.add_run(f"WEIGHT: {weight_total:.0f} KGS\n").bold = True
            extra_info_p.add_run(f"{container_count}×{container_type}").bold = True

            for r in extra_info_p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(16)
                r.bold = True

            bank_text = self.bank_details or (
                "Bank Name: EXPORT DEVELOPMENT BANK OF EGYPT\n"
                "Account Holder's Name: ELRAEE FOR DEHYDRATION\n"
                "Account Number: 2020273697801029\n"
                "IBAN: EG7400615611020273697801029\n"
                "Swift Code: EXDEEGCX"
            )

            doc.add_paragraph("")
            doc.add_paragraph("")
            for i, line in enumerate(bank_text.splitlines()):
                p = doc.add_paragraph(line)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.bold = True
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(8)

            doc.save(path)
            if preview:
                try:
                    os.startfile(path)
                except Exception:
                    pass
            else:
                QMessageBox.information(self, "Done", f"Saved Word: {os.path.basename(path)}")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to create Word:\n{e}")
